from app.repositories import aiRepo, fileRepo
from app.models.ai import ChatBody, ChatBodyV2
from app.models.file import NewFileRequest
from app.models.tables.databaseTables import Chat_Session, User_Model_Cfg, File
from datetime import datetime
from app.utils.chatbot import Chatbot
from app.utils.chatbotv2 import bots
import uuid
import json
import base64
import binascii
import re
from PIL import Image
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


NORMAL_EN_FONT = "Times New Roman"
NORMAL_CN_FONT = "宋体"
CODE_FONT = "Consolas"
USER_HEADING_COLOR = RGBColor(31, 78, 121)
ASSISTANT_HEADING_COLOR = RGBColor(83, 129, 53)
DEFAULT_HEADING_COLOR = RGBColor(64, 64, 64)


def _extract_cfg_items(payload):
    if isinstance(payload, dict):
        data = payload.get("data", [])
        return data if isinstance(data, list) else []
    if isinstance(payload, list):
        return payload
    return []


def _merge_user_cfg(kwargs_items, user_cfg_items):
    value_map = {}
    for cfg_item in user_cfg_items:
        if isinstance(cfg_item, dict) and "name" in cfg_item:
            value_map[cfg_item["name"]] = cfg_item.get("value")

    merged_items = []
    for kwargs_item in kwargs_items:
        if not isinstance(kwargs_item, dict):
            continue

        merged_item = dict(kwargs_item)
        item_name = merged_item.get("name")
        if item_name in value_map:
            merged_item["default"] = value_map[item_name]
        merged_items.append(merged_item)

    return merged_items


def _normalize_image_inputs(image):
    if image is None:
        raise ValueError("image is required")

    image_list = image if isinstance(image, list) else [image]
    normalized_images = []

    for item in image_list:
        if isinstance(item, (bytes, bytearray)):
            normalized_images.append(bytes(item))
            continue

        if not isinstance(item, str):
            raise ValueError("image items must be bytes or base64 string")

        encoded = item.strip()
        if "," in encoded and encoded.lower().startswith("data:image"):
            encoded = encoded.split(",", 1)[1]

        try:
            normalized_images.append(base64.b64decode(encoded, validate=True))
        except (binascii.Error, ValueError):
            raise ValueError("invalid base64 image data")

    return normalized_images


def _build_export_filename(session_name: str | None):
    base_name = (session_name or "session").strip()
    if not base_name:
        base_name = "session"
    base_name = base_name[:20]
    base_name = re.sub(r'[\\/:*?"<>|\r\n\t]', "_", base_name)
    return f"{base_name}.docx"


def _extract_image_ids(content: str):
    image_ids = []
    for line in content.splitlines():
        line = line.strip()
        if not line.startswith("Image-:"):
            continue
        image_id = line.replace("Image-:", "", 1).strip()
        if image_id:
            image_ids.append(image_id)
    return image_ids


def _decode_file_data(file_data: bytes | bytearray | memoryview | str):
    if not file_data:
        raise ValueError("empty file data")

    if isinstance(file_data, str):
        encoded = file_data.strip()
    elif isinstance(file_data, memoryview):
        encoded = file_data.tobytes().decode("utf-8").strip()
    else:
        encoded = bytes(file_data).decode("utf-8").strip()

    if "," in encoded and encoded.lower().startswith("data:"):
        encoded = encoded.split(",", 1)[1]
    return base64.b64decode(encoded)


def _split_markdown_table_row(line: str):
    row = line.strip()
    if row.startswith("|"):
        row = row[1:]
    if row.endswith("|"):
        row = row[:-1]
    return [cell.strip() for cell in row.split("|")]


def _is_markdown_table_separator(line: str):
    stripped = line.strip()
    return bool(re.match(r"^\|?\s*:?-{3,}:?(\s*\|\s*:?-{3,}:?)+\s*\|?$", stripped))


def _set_run_font(run, en_font: str = NORMAL_EN_FONT, cn_font: str = NORMAL_CN_FONT):
    run.font.name = en_font
    run_properties = run._r.get_or_add_rPr()
    fonts = run_properties.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        run_properties.append(fonts)
    fonts.set(qn("w:ascii"), en_font)
    fonts.set(qn("w:hAnsi"), en_font)
    fonts.set(qn("w:eastAsia"), cn_font)


def _set_document_default_fonts(document):
    # 设置文档默认中英文字体，避免 Word 对中文回退到非预期字体。
    normal_style = document.styles["Normal"]
    normal_style.font.name = NORMAL_EN_FONT
    normal_style.font.size = Pt(11)
    normal_style._element.rPr.rFonts.set(qn("w:ascii"), NORMAL_EN_FONT)
    normal_style._element.rPr.rFonts.set(qn("w:hAnsi"), NORMAL_EN_FONT)
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), NORMAL_CN_FONT)


def _set_run_shading(run, fill: str):
    run_properties = run._r.get_or_add_rPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    run_properties.append(shading)


def _set_paragraph_shading(paragraph, fill: str):
    paragraph_properties = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    paragraph_properties.append(shading)


def _set_cell_shading(cell, fill: str):
    cell_properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell_properties.append(shading)


def _set_table_borders(table):
    # 使用显式 XML 边框，确保不同 Word 客户端打开时表格边框稳定显示。
    table.style = "Table Grid"
    table_properties = table._tbl.tblPr
    borders = table_properties.find(qn("w:tblBorders"))
    if borders is not None:
        table_properties.remove(borders)

    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "8")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "808080")
        borders.append(border)
    table_properties.append(borders)


def _append_inline_text(paragraph, text: str):
    # 渲染 Markdown 行内样式，并将 <br> / <br/> 转换为 Word 段内换行。
    inline_pattern = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*|__[^_]+__|\*[^*\n]+\*|_[^_\n]+_)")
    position = 0

    def add_text_runs(run_text: str):
        runs = []
        parts = re.split(r"<br\s*/?>", run_text, flags=re.IGNORECASE)
        for part_index, part in enumerate(parts):
            if part_index > 0:
                paragraph.add_run().add_break()
            if not part:
                continue
            run = paragraph.add_run(part)
            _set_run_font(run)
            runs.append(run)
        return runs

    for match in inline_pattern.finditer(text):
        if match.start() > position:
            add_text_runs(text[position : match.start()])

        token = match.group(0)
        run_text = token
        is_bold = False
        is_italic = False
        is_code = False

        if token.startswith("`") and token.endswith("`"):
            run_text = token[1:-1]
            is_code = True
        elif token.startswith("**") and token.endswith("**"):
            run_text = token[2:-2]
            is_bold = True
        elif token.startswith("__") and token.endswith("__"):
            run_text = token[2:-2]
            is_bold = True
        elif token.startswith("*") and token.endswith("*"):
            run_text = token[1:-1]
            is_italic = True
        elif token.startswith("_") and token.endswith("_"):
            run_text = token[1:-1]
            is_italic = True

        runs = add_text_runs(run_text)
        if not runs:
            position = match.end()
            continue
        for run in runs:
            run.bold = is_bold
            run.italic = is_italic
            if is_code:
                _set_run_font(run, CODE_FONT, CODE_FONT)
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(196, 46, 46)
                _set_run_shading(run, "F6F8FA")

        position = match.end()

    if position < len(text):
        add_text_runs(text[position:])


def _append_code_block(document, code_lines: list[str]):
    # 将 Markdown 围栏代码块写成浅底色、Consolas 字体的独立段落。
    paragraph = document.add_paragraph()
    _set_paragraph_shading(paragraph, "F6F8FA")
    paragraph.paragraph_format.left_indent = Inches(0.15)
    paragraph.paragraph_format.right_indent = Inches(0.15)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run("\n".join(code_lines))
    _set_run_font(run, CODE_FONT, CODE_FONT)
    run.font.size = Pt(9)


def _append_table_cell_text(cell, text: str, bold: bool = False):
    # 表格单元格复用行内渲染，因此支持粗体、斜体、行内代码和 <br> 换行。
    paragraph = cell.paragraphs[0]
    _append_inline_text(paragraph, text)
    for run in paragraph.runs:
        run.bold = bold or run.bold


def _append_role_heading(document, role: str, fallback: str):
    # User / Assistant 作为一级标题输出，并用颜色区分对话身份。
    role_name = role.capitalize() if role else fallback
    heading = document.add_heading(level=1)
    run = heading.add_run(role_name)
    _set_run_font(run)
    run.bold = True
    if role.lower() == "user":
        run.font.color.rgb = USER_HEADING_COLOR
    elif role.lower() == "assistant":
        run.font.color.rgb = ASSISTANT_HEADING_COLOR
    else:
        run.font.color.rgb = DEFAULT_HEADING_COLOR
    return heading


def _append_markdown(document, md_content: str):
    # 将会话 Markdown 内容渲染为 Word 结构；正文标题自动下移一级，避开角色标题。
    lines = md_content.splitlines()
    index = 0

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()

        if stripped.startswith("```"):
            code_lines = []
            index += 1
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code_lines.append(lines[index])
                index += 1
            if index < len(lines):
                index += 1
            _append_code_block(document, code_lines)
            continue

        if not stripped:
            index += 1
            continue

        if stripped.startswith("#"):
            md_level = len(stripped) - len(stripped.lstrip("#"))
            level = min(md_level + 1, 9)
            heading_text = stripped[md_level:].strip()
            if heading_text:
                heading = document.add_heading(level=level)
                _append_inline_text(heading, heading_text)
                for run in heading.runs:
                    run.font.color.rgb = DEFAULT_HEADING_COLOR
            index += 1
            continue

        if index + 1 < len(lines) and "|" in lines[index] and _is_markdown_table_separator(lines[index + 1]):
            headers = _split_markdown_table_row(lines[index])
            table = document.add_table(rows=1, cols=len(headers))
            _set_table_borders(table)
            for col_index, header in enumerate(headers):
                cell = table.rows[0].cells[col_index]
                _set_cell_shading(cell, "D9EAF7")
                _append_table_cell_text(cell, header, bold=True)

            index += 2
            while index < len(lines):
                row_line = lines[index].strip()
                if not row_line or "|" not in row_line:
                    break
                row_values = _split_markdown_table_row(lines[index])
                row_cells = table.add_row().cells
                for col_index in range(len(headers)):
                    cell_text = row_values[col_index] if col_index < len(row_values) else ""
                    _append_table_cell_text(row_cells[col_index], cell_text)
                index += 1
            continue

        if stripped.startswith("- ") or stripped.startswith("* "):
            paragraph = document.add_paragraph(style="List Bullet")
            _append_inline_text(paragraph, stripped[2:].strip())
            index += 1
            continue

        ordered_match = re.match(r"^\d+\.\s+(.*)$", stripped)
        if ordered_match:
            paragraph = document.add_paragraph(style="List Number")
            _append_inline_text(paragraph, ordered_match.group(1))
            index += 1
            continue

        paragraph_lines = [line]
        index += 1
        while index < len(lines):
            next_line = lines[index]
            next_stripped = next_line.strip()
            if not next_stripped:
                break
            if next_stripped.startswith("#"):
                break
            if next_stripped.startswith("- ") or next_stripped.startswith("* "):
                break
            if re.match(r"^\d+\.\s+", next_stripped):
                break
            if next_stripped.startswith("```"):
                break
            if index + 1 < len(lines) and "|" in next_line and _is_markdown_table_separator(lines[index + 1]):
                break
            paragraph_lines.append(next_line)
            index += 1

        paragraph = document.add_paragraph()
        _append_inline_text(paragraph, "\n".join(paragraph_lines).strip())


async def export_session_to_word(session_id: uuid.UUID, user_id: uuid.UUID):
    # 组装会话 Word 导出：解析消息、批量加载图片、渲染 Markdown，最后返回 docx 二进制。
    session = aiRepo.select_session_by_id_and_user_id(session_id, user_id)
    if session is None:
        return None

    messages = []
    content = session.content if session.content else {}
    if isinstance(content, dict):
        messages = content.get("messages", [])
    if not isinstance(messages, list):
        messages = []

    image_uuid_list = []
    for message in messages:
        if not isinstance(message, dict):
            continue
        message_content = message.get("content", "")
        if not isinstance(message_content, str):
            continue
        image_ids = _extract_image_ids(message_content)
        for image_id in image_ids:
            try:
                image_uuid_list.append(uuid.UUID(image_id))
            except ValueError:
                continue

    image_file_map = {str(file.id): file for file in fileRepo.select_files_by_ids(image_uuid_list)}  # type: ignore

    document = Document()
    _set_document_default_fonts(document)
    title = document.add_heading(session.session_name or "Session", level=0)
    for run in title.runs:
        _set_run_font(run)

    for index, message in enumerate(messages, start=1):
        if not isinstance(message, dict):
            continue

        role = str(message.get("role", "")).strip()
        message_content = message.get("content", "")
        if not isinstance(message_content, str):
            message_content = str(message_content)

        _append_role_heading(document, role, f"Message {index}")

        image_ids = _extract_image_ids(message_content)
        if image_ids:
            for image_id in image_ids:
                image_file = image_file_map.get(image_id)
                if image_file is None:
                    document.add_paragraph(f"[Image not found: {image_id}]")
                    continue
                try:
                    image_bytes = _decode_file_data(image_file.data)
                    document.add_picture(BytesIO(image_bytes), width=Inches(6))
                except Exception:
                    document.add_paragraph(f"[Image decode failed: {image_id}]")
            continue

        _append_markdown(document, message_content)

    output = BytesIO()
    document.save(output)
    output.seek(0)
    return {
        "filename": _build_export_filename(session.session_name),
        "data": output.getvalue(),
    }


async def get_user_sessions(user_id):
    return aiRepo.select_sessions_by_user_id(user_id)


async def get_session_content(session_id):
    sessions = aiRepo.select_sessions_by_session_id(session_id)
    if sessions:
        return sessions
    else:
        return None


async def generate_session_name(user_input: str) -> str:
    return await Chatbot().generate_session_name(user_input)


async def add_session(chat_body: ChatBody, user_id):
    session_name = await generate_session_name(
        "user:"
        + chat_body.content.get("messages", [{}])[0].get("content", "新会话")  # type: ignore
        + "\n assistant:"
        + chat_body.content.get("messages", [{}])[1].get("content", "")[:1000]  # type: ignore
    )
    session = Chat_Session(
        user_id=user_id,
        session_name=session_name,
        create_time=chat_body.create_time or datetime.now(),
        content=chat_body.content or {},
    )
    return aiRepo.insert_chat_session(session)


async def update_session(chat: Chat_Session):
    return aiRepo.update_chat_session_content(chat.id, chat.content)


async def delete_session(session_id: uuid.UUID):
    return aiRepo.delete_chat_session(session_id)


async def get_models():
    return aiRepo.select_models()


async def get_models_v2(user_id: uuid.UUID):
    models = aiRepo.select_models_v2()
    user_cfg_map = aiRepo.select_user_model_cfg_map(user_id, [model.id for model in models])

    model_list = []
    for model in models:
        kwargs_items = _extract_cfg_items(model.kwargs)
        user_cfg_items = _extract_cfg_items(user_cfg_map.get(model.id))
        merged_kwargs = _merge_user_cfg(kwargs_items, user_cfg_items)

        model_list.append(
            {
                "modelType": model.model_type,
                "model": model.model,
                "kwargs": merged_kwargs,
            }
        )

    return model_list


async def chat_stream(model: str, messages: dict, **kwargs):
    model_data = aiRepo.select_model_v2_by_model(model)
    if model not in bots or model_data is None:
        raise ValueError(f"Model '{model}' not found in database.")
    bot = bots[model]
    if model_data.sdk_type == "openai" or model_data.sdk_type == "volcengine":
        response = await bot.chat.completions.create(
            model=model,
            messages=messages,
            stream=True,
            **kwargs,
        )
        async for chunk in response:
            if not chunk.choices:
                continue
            delta = chunk.choices[0].delta
            content = getattr(delta, "content", None)
            reasoning = getattr(delta, "reasoning_content", None)
            if content:
                yield f"data: {json.dumps({'content': content}, ensure_ascii=False)}\n\n"
            if reasoning:
                yield f"data: {json.dumps({'reasoning_content': reasoning}, ensure_ascii=False)}\n\n"
        yield "data: [DONE]\n\n"


async def save_file(file_req: NewFileRequest):
    file = File(
        source_url=file_req.source_url,
        filename=file_req.filename,
        file_type=file_req.file_type,
        data=file_req.data,
    )
    return fileRepo.insert_file(file)


async def get_file_by_id(file_id: uuid.UUID):
    return fileRepo.select_file_by_id(file_id)


# 将base64形式的图像数据解码为字节，并使用PIL库进行压缩，最后再编码回base64格式
async def compress_file_data(data: bytes) -> bytes:
    img = base64.b64decode(data)
    img = Image.open(BytesIO(img))
    if img.mode not in ("RGB", "L"):
        img = img.convert("RGB")
    # 在内存中压缩，不落盘
    buffer = BytesIO()
    img.save(buffer, format="JPEG", quality=90, optimize=True)
    compressed_image_data = buffer.getvalue()
    compressed_image_base64 = base64.b64encode(compressed_image_data)
    return compressed_image_base64


async def image_generate(model: str, prompt: str, **kwargs):
    model_data = aiRepo.select_model_v2_by_model(model)
    if model not in bots or model_data is None:
        raise ValueError(f"Model '{model}' not found in database.")
    bot = bots[model]
    response = await bot.images.generate(
        model=model,
        prompt=prompt,
        **kwargs,
    )
    images = []
    for img in response.data:
        if img.b64_json:
            images.append({"type": "b64_json", "data": img.b64_json})
        elif img.url:
            images.append({"type": "url", "data": img.url})
    return images


async def image_edit(model: str, image: list, prompt: str, **kwargs):
    model_data = aiRepo.select_model_v2_by_model(model)
    if model not in bots or model_data is None:
        raise ValueError(f"Model '{model}' not found in database.")
    bot = bots[model]
    normalized_images = _normalize_image_inputs(image)
    # print("开始调用图像编辑接口")
    # with open("debug_input_image.png", "wb") as f:
    #     f.write(normalized_images[0])
    response = await bot.images.edit(
        model=model,
        image=normalized_images,
        prompt=prompt,
        **kwargs,
    )
    images = []
    for img in response.data:
        if img.b64_json:
            images.append({"type": "b64_json", "data": img.b64_json})
        elif img.url:
            images.append({"type": "url", "data": img.url})
    return images


async def update_user_model_cfg(user_id: uuid.UUID, chat_body: ChatBodyV2):
    modelv2 = aiRepo.select_model_v2_by_model(chat_body.model)
    if modelv2 is None:
        raise ValueError(f"Model '{chat_body.model}' not found in database.")
    if chat_body.kwargs is None or chat_body.kwargs == {}:
        return
    cfg_row = aiRepo.select_user_model_cfg(user_id, modelv2.id)

    cfg_payload = cfg_row.cfg if cfg_row else {}
    existing_items = _extract_cfg_items(cfg_payload)

    value_map = {}
    for item in existing_items:
        if isinstance(item, dict) and "name" in item:
            value_map[item["name"]] = item.get("value")

    for key, value in chat_body.kwargs.items():
        value_map[key] = value

    merged_cfg = {"data": [{"name": key, "value": value} for key, value in value_map.items()]}

    if cfg_row:
        return aiRepo.update_user_model_cfg_cfg(cfg_row.id, merged_cfg)

    new_cfg = aiRepo.insert_user_model_cfg(
        user_model_cfg=User_Model_Cfg(
            user_id=user_id,
            model_v2_id=modelv2.id,
            cfg=merged_cfg,
        )
    )
    return new_cfg
